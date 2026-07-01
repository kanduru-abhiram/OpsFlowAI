from io import BytesIO
import re

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.services.openai_client import OpenAIService


def extract_text_from_upload(filename: str, content_type: str | None, raw: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    mime = (content_type or "").lower()
    if suffix in {"png", "jpg", "jpeg", "webp", "bmp", "gif", "tiff"} or mime.startswith("image/"):
        return OpenAIService().extract_image_text(filename=filename, content_type=content_type or "image/png", raw=raw)
    if suffix == "pdf" or "pdf" in mime:
        reader = PdfReader(BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if suffix == "docx" or "wordprocessingml" in mime:
        document = DocxDocument(BytesIO(raw))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        return "\n".join(paragraphs + table_cells).strip()
    return raw.decode("utf-8", errors="ignore")


def extract_entities(text: str) -> dict[str, str]:
    email = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    account = re.search(r"\b(?:AC|ACC|Account)[:\-\s]*([A-Z0-9-]{6,})", text, re.IGNORECASE)
    ref = re.search(r"\b(?:REF|SR|TKT)[:\-\s]*([A-Z0-9-]{5,})", text, re.IGNORECASE)
    name = re.search(r"(?:Customer|Name)[:\-\s]*([A-Z][a-z]+(?:\s[A-Z][a-z]+){0,2})", text)
    return {
        "email": email.group(0) if email else "",
        "account_number": account.group(1) if account else "",
        "reference_id": ref.group(1) if ref else "",
        "customer_name": name.group(1) if name else "",
    }


def summarize_document(text: str) -> str:
    clean = " ".join(text.split())
    if len(clean) <= 220:
        return clean
    return clean[:220].rsplit(" ", 1)[0] + "..."
